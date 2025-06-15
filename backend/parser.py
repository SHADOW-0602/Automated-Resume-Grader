import re
import pdfplumber
from pdf2image import convert_from_bytes
import pytesseract
import nltk
from dateutil import parser as date_parser
from datetime import datetime
import fitz  # PyMuPDF
import docx
import logging
import spacy
from langdetect import detect
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTFigure
import warnings
import cv2
import numpy as np

# Setup logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Suppress warnings from pdfminer
warnings.filterwarnings("ignore", category=UserWarning)

# Load models
try:
    nlp = spacy.load("en_core_web_lg")
except OSError:
    logger.error("Spacy model 'en_core_web_lg' not found. Please install it.")
    nlp = None

# Download NLTK data
nltk.download('punkt', quiet=True)
nltk.download('maxent_ne_chunker', quiet=True)
nltk.download('words', quiet=True)

def detect_resume_language(text):
    """Detect the primary language of the resume text."""
    try:
        return detect(text[:1000]) if text else 'en'
    except:
        return 'en'

def preprocess_image_for_ocr(image, language='en'):
    """Preprocess image for better OCR results with language support."""
    image = image.convert('L')  # Grayscale
    image = cv2.adaptiveThreshold(np.array(image), 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
    return image

def is_scanned_pdf(file_obj):
    """Check if PDF is scanned using text extraction confidence."""
    try:
        file_obj.seek(0)
        doc = fitz.open(stream=file_obj.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        
        file_obj.seek(0)
        with pdfplumber.open(file_obj) as pdf:
            total_confidence = 0
            char_count = 0
            for page in pdf.pages:
                if page.chars:
                    confidences = [char.get('confidence', 0) for char in page.chars]
                    if confidences:
                        total_confidence += sum(confidences)
                        char_count += len(confidences)
        
        avg_confidence = (total_confidence / char_count) if char_count > 0 else 0
        text_length = len(text.strip())
        
        logger.debug(f"Scanned PDF check: Text length = {text_length}, Avg confidence = {avg_confidence:.2f}")
        return text_length < 100 and avg_confidence < 70
    except Exception as e:
        logger.error(f"Error checking scanned PDF: {str(e)}")
        return True

def extract_text_with_layout(pdf_path):
    """Extract text while preserving layout structure using pdfminer."""
    text = ""
    for page_layout in extract_pages(pdf_path):
        for element in page_layout:
            if isinstance(element, LTTextContainer):
                text += element.get_text() + "\n"
            elif isinstance(element, LTFigure):
                # Handle figures (potential text in images)
                text += "[IMAGE]"  # Placeholder for images
    return text

def extract_text_from_pdf(file_obj):
    """Extract text from PDF with improved layout handling and OCR."""
    try:
        file_obj.seek(0)
        if is_scanned_pdf(file_obj):
            logger.info("Processing as scanned PDF")
            language = detect_resume_language("")
            file_obj.seek(0)
            images = convert_from_bytes(file_obj.read())
            text = ""
            for image in images:
                preprocessed = preprocess_image_for_ocr(image, language)
                page_text = pytesseract.image_to_string(
                    preprocessed,
                    lang=language if language in ['en', 'fr', 'de', 'es'] else 'eng'
                )
                text += page_text + "\n"
            logger.debug(f"OCR extracted text length: {len(text)}")
        else:
            logger.info("Processing as text-based PDF with layout preservation")
            file_obj.seek(0)
            with pdfplumber.open(file_obj) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text(layout=True) or ""
                    text += page_text + "\n"
            logger.debug(f"Extracted text length: {len(text)}")
        
        ats_warnings = check_ats_compliance(text)
        if ats_warnings:
            logger.warning(f"ATS compliance issues: {ats_warnings}")
        
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file_obj):
    """Extract text from DOCX with formatting awareness."""
    try:
        file_obj.seek(0)
        doc = docx.Document(file_obj)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        if doc.tables:
            logger.warning("DOCX contains tables which may affect ATS parsing")
            text += "\n[TABLE CONTENT DETECTED]"
        
        logger.debug(f"DOCX extracted text length: {len(text)}")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_text_from_txt(file_obj):
    """Extract text from TXT file with encoding detection."""
    try:
        file_obj.seek(0)
        text = file_obj.read().decode('utf-8', errors='ignore')
        logger.debug(f"TXT extracted text length: {len(text)}")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        return ""

def extract_name(text):
    """Extract name using both Spacy and NLTK for better accuracy."""
    if nlp:
        try:
            doc = nlp(text[:1000])  # Only process first part for efficiency
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    logger.debug(f"Spacy extracted name: {ent.text}")
                    return ent.text
        except Exception as e:
            logger.warning(f"Spacy name extraction failed: {str(e)}")
    
    # Fallback to NLTK
    try:
        sentences = nltk.sent_tokenize(text)
        for sent in sentences[:5]:
            tokens = nltk.word_tokenize(sent)
            tagged = nltk.pos_tag(tokens)
            entities = nltk.chunk.ne_chunk(tagged)
            for subtree in entities:
                if isinstance(subtree, nltk.Tree) and subtree.label() == 'PERSON':
                    name = ' '.join(word for word, pos in subtree.leaves())
                    logger.debug(f"NLTK extracted name: {name}")
                    return name
    except Exception as e:
        logger.warning(f"NLTK name extraction failed: {str(e)}")
    
    # Final regex fallback
    name_match = re.search(r'^([A-Z][a-z]+(?: [A-Z][a-z]+)+)', text, re.MULTILINE)
    if name_match:
        logger.debug(f"Regex extracted name: {name_match.group(1)}")
        return name_match.group(1)
    
    logger.debug("No name extracted, defaulting to 'Unknown Name'")
    return "Unknown Name"

def check_ats_compliance(text):
    """Enhanced ATS compliance checks."""
    warnings = []
    
    # Check for problematic elements
    if re.search(r'\b(table|header|footer|text box)\b', text, re.IGNORECASE):
        warnings.append("Avoid tables/headers/footers/text boxes for better ATS parsing.")
    
    # Check length
    word_count = len(text.split())
    if word_count > 1000:
        warnings.append(f"Resume is too long ({word_count} words). Ideal: 400-800 words.")
    elif word_count < 200:
        warnings.append(f"Resume is very short ({word_count} words). Consider adding more details.")
    
    # Check for graphics
    if "[IMAGE]" in text:
        warnings.append("Images detected which may not be parsed by ATS.")
    
    # Check for unusual formatting
    if re.search(r'[●♦♣♥♠]', text):  # Graphic bullets
        warnings.append("Unusual bullet points detected that may not parse well.")
    
    return warnings

def parse_resume(file_obj, job_keywords=None):
    """Parse resume with enhanced features."""
    try:
        extension = file_obj.filename.rsplit('.', 1)[1].lower() if hasattr(file_obj, 'filename') else 'txt'
        logger.debug(f"Parsing file with extension: {extension}")
        
        # Extract text based on file type
        if extension == 'pdf':
            text = extract_text_from_pdf(file_obj)
        elif extension == 'docx':
            text = extract_text_from_docx(file_obj)
        elif extension == 'txt':
            text = extract_text_from_txt(file_obj)
        else:
            logger.error(f"Unsupported file extension: {extension}")
            return empty_resume()
        
        if not text.strip():
            logger.error("No text extracted from file")
            return empty_resume()
        
        # Detect language
        language = detect_resume_language(text)
        logger.debug(f"Detected resume language: {language}")
        
        # Extract contact info
        contact = {
            "name": extract_name(text),
            "email": extract_email(text),
            "phone": extract_phone(text),
            "linkedin": extract_linkedin(text),
            "language": language
        }
        
        # Extract sections
        sections = {
            "summary": extract_summary(text),
            "experience": extract_experience(text),
            "education": extract_education(text),
            "skills": extract_skills(text),
            "certifications": extract_certifications(text),
            "ats_warnings": check_ats_compliance(text)
        }

        # Calculate experience years
        experience_years = calculate_experience_years(sections["experience"])
        seniority = detect_role_seniority(text, experience_years)
        
        # Determine seniority before using it in metrics
        seniority = detect_role_seniority(text)
        
        # Additional metrics
        metrics = {
            "experience_years": experience_years,
            "skill_count": len(sections["skills"]),
            "quantified_achievements": count_quantified_achievements(sections["experience"]),
            "action_verbs": count_action_verbs(text),
            "word_count": len(text.split()),
            "seniority": seniority,
            "section_weights": {
                "experience": 0.6 if seniority in ["Senior Leader", "Manager"] else 0.5 if seniority == "Senior" else 0.4,
                "skills": 0.3 if seniority == "Junior" else 0.2,
                "education": 0.2 if seniority == "Junior" else 0.1,
                "certifications": 0.1
            }
        }
        
        resume_data = {
            "raw_text": text,
            "contact": contact,
            "sections": sections,
            "metrics": metrics
        }
        
        logger.info("Resume parsing completed successfully")
        return resume_data
        
    except Exception as e:
        logger.error(f"Parsing failed: {str(e)}")
        return empty_resume()
    
def detect_role_seniority(text, experience_years=0):
    seniority_keywords = {
        "Junior": ["assistant", "junior", "entry-level"],
        "Senior": ["senior", "lead", "principal"],
        "Manager": ["manager", "director", "head of", "vp", "vice president", "chief"]
    }
    leadership_titles = ["manager", "director", "head of", "vp", "chief", "leader", "supervisor"]
    
    # Check for leadership keywords in text
    is_leader = any(re.search(rf'\b{kw}\b', text, re.IGNORECASE) for kw in leadership_titles)
    
    # Determine seniority
    for level, keywords in seniority_keywords.items():
        if any(re.search(rf'\b{kw}\b', text, re.IGNORECASE) for kw in keywords):
            if level == "Manager" and experience_years >= 8 and is_leader:
                return "Senior Leader"
            return level
    
    # Default to Mid-Level if no keywords match
    return "Mid-Level"   

def empty_resume():
    """Return empty resume structure."""
    logger.debug("Returning empty resume structure")
    return {
        "raw_text": "",
        "contact": {"name": "Unknown Name", "email": "", "phone": "", "linkedin": ""},
        "sections": {
            "summary": "",
            "experience": [],
            "education": [],
            "skills": [],
            "certifications": []
        },
        "metrics": {
            "experience_years": 0,
            "skill_count": 0,
            "quantified_achievements": 0,
            "action_verbs": 0
        },
        "word_count": 0
    }

def extract_name(text):
    """Extract name using NLTK for NER and regex fallback."""
    try:
        sentences = nltk.sent_tokenize(text)
        for sent in sentences[:5]:
            tokens = nltk.word_tokenize(sent)
            tagged = nltk.pos_tag(tokens)
            entities = nltk.chunk.ne_chunk(tagged)
            for subtree in entities:
                if isinstance(subtree, nltk.Tree) and subtree.label() == 'PERSON':
                    name = ' '.join(word for word, pos in subtree.leaves())
                    logger.debug(f"Extracted name via NLTK: {name}")
                    return name
    except Exception as e:
        logger.warning(f"NLTK name extraction failed: {str(e)}")
    
    name_match = re.search(r'^((?:Dr\.|Mr\.|Ms\.|Mrs\.|Prof\.)?\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+(?:\s+[A-Z]+)?)', text, re.MULTILINE)
    if name_match:
        logger.debug(f"Extracted name via regex: {name_match.group(1)}")
        return name_match.group(1)
    
    logger.debug("No name extracted, defaulting to 'Unknown Name'")
    return "Unknown Name"

def extract_email(text):
    """Extract email with improved regex."""
    email = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email:
        logger.debug(f"Extracted email: {email.group(0)}")
        return email.group(0)
    logger.debug("No email extracted")
    return ""

def extract_phone(text):
    """Extract phone number with international support."""
    phone = re.search(r'(\+?\d{1,3}[-.\s]?)?(?:\(\d{3}\)|\d{3})[-.\s]?\d{3}[-.\s]?\d{4}(?:\s*(?:ext\.?|x)\s*\d+)?', text)
    if phone:
        logger.debug(f"Extracted phone: {phone.group(0)}")
        return phone.group(0)
    logger.debug("No phone extracted")
    return ""

def extract_linkedin(text):
    """Extract LinkedIn URL."""
    linkedin = re.search(r'(https?://)?(?:www\.)?linkedin\.com/(in|company)/[a-zA-Z0-9-]+/?', text)
    if linkedin:
        logger.debug(f"Extracted LinkedIn: {linkedin.group(0)}")
        return linkedin.group(0)
    logger.debug("No LinkedIn extracted")
    return ""

def extract_summary(text):
    """Extract summary section with case-insensitive keywords."""
    summary_keywords = [r'summary', r'profile', r'about', r'objective', r'overview']
    for keyword in summary_keywords:
        pattern = rf'\b{keyword}\b.*?(?=\n\n|\Z)'
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if match:
            summary = match.group(0).strip()
            if len(summary) > 20:
                logger.debug(f"Extracted summary: {summary[:100]}...")
                return summary
    first_para = text.split('\n\n')[0].strip()
    if len(first_para.split()) > 20 and not any(kw in first_para.lower() for kw in ['experience', 'education', 'skills']):
        logger.debug(f"Extracted summary via fallback: {first_para[:100]}...")
        return first_para
    logger.debug("No summary extracted")
    return ""

def extract_education(text):
    """Extract education with case-insensitive keywords and header-less content."""
    education = []
    education_keywords = [
        r'education', r'academic background', r'qualifications', r'degree',
        r'academic history', r'training', r'certifications', r'certificates'
    ]
    
    for keyword in education_keywords:
        pattern = rf'\b{keyword}\b.*?(?=\n\n|\Z)'
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        for match in matches:
            education_section = match.group(0)
            degrees = re.findall(
                r'(?:B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|Bachelor|Master|Doctorate|Certificate|Certification)[\w\s,.-]*(?:\d{4})?(?:.*?University|Institute|College|Academy)?',
                education_section, re.IGNORECASE
            )
            certs = re.findall(r'CompTIA\s+\w+|\b[A-Z]{2,}\s+Certification\b|Certified\s+\w+', education_section, re.IGNORECASE)
            education.extend([e.strip() for e in degrees + certs if e.strip()])
    
    global_edu = re.findall(
        r'(?:B\.?S\.?|B\.?A\.?|M\.?s\.?|M\.?A\.?|Ph\.?D\.?|Bachelor|Master|Doctorate|Certificate|Certification)[\w\s,.-]*(?:\d{4})?(?:.*?University|Institute|College|Academy)?',
        text, re.IGNORECASE
    )
    global_certs = re.findall(r'CompTIA\s+\w+|\b[A-Z]{2,}\s+Certification\b|Certified\s+\w+', text, re.IGNORECASE)
    education.extend([e.strip() for e in global_edu + global_certs if e.strip()])
    
    try:
        sentences = nltk.sent_tokenize(text)
        for sent in sentences:
            tokens = nltk.word_tokenize(sent)
            tagged = nltk.pos_tag(tokens)
            entities = nltk.chunk.ne_chunk(tagged)
            for subtree in entities:
                if isinstance(subtree, nltk.Tree) and subtree.label() == 'ORGANIZATION':
                    org = ' '.join(word for word, pos in subtree.leaves())
                    if re.search(r'University|Institute|College|Academy', org, re.IGNORECASE):
                        education.append(org.strip())
    except Exception as e:
        logger.warning(f"NLTK education extraction failed: {str(e)}")
    
    education = list(set(education))
    logger.debug(f"Extracted education: {education[:3]}")
    return education

def score_project_impact(text):
    impact_terms = [
        r'led \d+ team', r'\$\d+ budget', r'enterprise-wide',
        r'cross-functional', r'multi-year'
    ]
    return sum(1 for term in impact_terms if re.search(term, text, re.IGNORECASE))

def extract_skills(text):
    """Extract skills with case-insensitive keywords and header-less content."""
    skills = []
    skill_keywords = [
        'python', 'java', 'sql', 'javascript', 'machine learning', 'aws', 'docker',
        'active directory', 'powershell', 'vmware', 'office 365', 'azure', 'networking',
        'linux', 'exchange', 'backup', 'disaster recovery', 'html', 'css', 'c++', 'bash',
        't-sql', 'group policy', 'storage management', 'cisco', 'firewall', 'vpn', 'epo',
        'sharepoint', 'sccm', 'netbackup', 'data domain', 'mcafee', 'solarwinds',
        'storsimple', 'twinstrata', 'enterprise vault', 'adfs', 'dfs', 'hyper-v', 'iaas',
        'scsi', 'virtual desktop infrastructure', 'isa server', 'typescript', 'node.js',
        'react', 'angular', 'c#', 'go', 'rust', 'nosql', 'mongodb', 'postgresql', 'mysql',
        'graphql', 'kubernetes', 'docker swarm', 'terraform', 'ansible', 'cloudformation',
        'palo alto', 'fortinet', 'siem', 'splunk', 'wireshark', 'penetration testing',
        'encryption', 'oauth', 'saml', 'ids/ips', 'jira', 'confluence', 'servicenow',
        'nagios', 'zabbix', 'prometheus', 'grafana', 'citrix', 'vmware esxi', 'aws ec2',
        'aws s3', 'google cloud platform', 'gcp'
    ]
    
    skills_keywords = [r'skills', r'highlights', r'technical skills', r'competencies', r'proficiencies', r'abilities']
    for keyword in skills_keywords:
        pattern = rf'\b{keyword}\b.*?(?=\n\n|\Z)'
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        for match in matches:
            skills_section = match.group(0)
            skills.extend(re.findall(r'(?:•|-|\*|,)\s*([^\n,]+)(?=(?:,|\n|$))', skills_section, re.IGNORECASE))
            for kw in skill_keywords:
                if re.search(rf'\b{kw}\b', skills_section, re.IGNORECASE):
                    skills.append(kw)
    
    for kw in skill_keywords:
        if re.search(rf'\b{kw}\b', text, re.IGNORECASE):
            skills.append(kw)
    
    extra_skills = re.findall(
        r'\b(?:Python|Java|SQL|JavaScript|C\+\+|HTML|CSS|Bash|T-SQL|PowerShell|VBScript|VMware|Azure|AWS|Docker|Linux|Windows Server|Active Directory|Exchange|Office 365|SharePoint|SCCM|NetBackup|Data Domain|McAfee|SolarWinds|Cisco|TypeScript|Node\.js|React|Angular|C#|Go|Rust|MongoDB|PostgreSQL|MySQL|GraphQL|Kubernetes|Terraform|Ansible|Palo Alto|Fortinet|Splunk|ServiceNow|Nagios|Zabbix|Prometheus|Grafana|Citrix)\b',
        text, re.IGNORECASE
    )
    skills.extend(extra_skills)
    
    skills = list(set([s.strip() for s in skills if s.strip()]))
    logger.debug(f"Extracted skills: {skills[:10]}")
    return skills

def extract_experience(text):
    """Extract experience with case-insensitive keywords and header-less content."""
    experience = []
    exp_keywords = [
        r'experience', r'work history', r'employment', r'professional experience',
        r'career history', r'job history', r'positions held', r'work experience'
    ]
    
    for keyword in exp_keywords:
        pattern = rf'\b{keyword}\b.*?(?=\n\n|\Z)'
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        for match in matches:
            exp_section = match.group(0)
            jobs = re.split(
                r'\n(?=(?:[A-Z][a-z]+ )*(?:Technician|Engineer|Administrator|Manager|Specialist|Consultant|Analyst|Developer|Liaison)\b.*?,\s*(?:[A-Z][a-z]+,)?\s*(?:\d{4}|Present|present|Current|current)(?:\s*(?:[-–]\s*(?:\d{4}|Present|present|Current|current))?))',
                exp_section, re.IGNORECASE
            )
            for job in jobs:
                if job.strip() and len(job.strip()) > 20 and re.search(r'[A-Z][a-z]+', job):
                    experience.append(job.strip())
    
    job_patterns = re.findall(
        r'(?:[A-Z][a-z]+ )*(?:Technician|Engineer|Administrator|Manager|Specialist|Consultant|Analyst|Developer|Liaison)\b.*?(?:\d{4}|Present|present|Current|current)(?:\s*(?:[-–]\s*(?:\d{4}|Present|present|Current|current))?)(?=\n\n|\Z)',
        text, re.IGNORECASE | re.DOTALL
    )
    for job in job_patterns:
        if job.strip() and len(job.strip()) > 10:
            experience.append(job.strip())
    
    experience = list(set(experience))
    logger.debug(f"Extracted experience: {experience[:2]}")
    return experience

def extract_certifications(text):
    """Extract certifications with validation."""
    cert_keywords = [r'certifications', r'certificates', r'credentials']
    certifications = []
    
    for keyword in cert_keywords:
        pattern = rf'\b{keyword}\b.*?(?=\n\n|\Z)'
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
        for match in matches:
            cert_section = match.group(0)
            certs = re.findall(r'([A-Za-z0-9]+(?:\s+[A-Za-z0-9\+#]+)*\s+(?:Certification|Certified|Certificate))', cert_section, re.IGNORECASE)
            certifications.extend([c.strip() for c in certs if c.strip()])
    
    global_certs = re.findall(r'([A-Za-z0-9]+(?:\s+[A-Za-z0-9\+#]+)*\s+(?:Certification|Certified|Certificate))', text, re.IGNORECASE)
    certifications.extend([c.strip() for c in global_certs if c.strip()])
    
    certifications = list(set(certifications))
    logger.debug(f"Extracted certifications: {certifications}")
    return certifications

def check_ats_compliance(text):
    warnings = []
    if re.search(r'[^\x00-\x7F]', text): 
        warnings.append("Non-ASCII characters detected, which may cause ATS parsing issues.")
    if re.search(r'\b(table|header|footer)\b', text, re.IGNORECASE):
        warnings.append("Avoid tables/headers/footers for better ATS parsing.")
    if len(text.split()) > 1000:
        warnings.append("Resume is too long (>1000 words). Ideal: 400-800 words.")
    return warnings

def calculate_experience_years(experience_items):
    """Calculate total years of experience from job entries."""
    total_years = 0
    date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}|\d{4}'
    
    for job in experience_items:
        dates = re.findall(date_pattern, job, re.IGNORECASE)
        if len(dates) >= 2:
            try:
                start_date = date_parser.parse(dates[0])
                end_date = date_parser.parse(dates[1]) if 'present' not in dates[1].lower() else datetime.now()
                total_years += (end_date - start_date).days / 365.25
            except Exception as e:
                logger.warning(f"Date parsing failed for job: {job[:50]}... Error: {str(e)}")
                continue
        elif 'present' in job.lower() or 'current' in job.lower():
            try:
                start_date = date_parser.parse(dates[0]) if dates else datetime.now()
                end_date = datetime.now()
                total_years += (end_date - start_date).days / 365.25
            except Exception as e:
                logger.warning(f"Date parsing failed for job: {job[:50]}... Error: {str(e)}")
                continue
    
    years = round(total_years, 1)
    logger.debug(f"Calculated experience years: {years}")
    return years

def count_quantified_achievements(experience_items):
    """Count quantified achievements in experience section."""
    count = 0
    quantifiers = [
        r'\d+%', r'\$\d+', r'\d+\+', r'\d+x', 
        r'increased by \d+', r'reduced by \d+', r'saved \$\d+',
        r'cut \w+ time by \d+', r'improved \w+ by \d+'
    ]
    
    for job in experience_items:
        for q in quantifiers:
            if re.search(q, job, re.IGNORECASE):
                count += 1
                break
    
    logger.debug(f"Quantified achievements count: {count}")
    return count

def count_action_verbs(text):
    """Count action verbs in text."""
    action_verbs = [
        'achieved', 'developed', 'implemented', 'led', 'managed', 'created',
        'designed', 'improved', 'increased', 'reduced', 'optimized', 'built'
    ]
    count = sum(text.lower().count(verb) for verb in action_verbs)
    logger.debug(f"Action verbs count: {count}")
    return count